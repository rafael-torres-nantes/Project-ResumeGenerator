"""Serviço para geração de currículo em DOCX e PDF."""

import json
import logging
from pathlib import Path
from typing import List, Tuple, Dict, Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


class CurriculumGenerator:
    """Gera o currículo em DOCX e converte para PDF.

    Configura margens, estilos e escreve as seções baseadas no arquivo JSON
    em personal_info/curriculum_data.json. Utiliza COM para converter para PDF.
    """

    _FONT_NAME = "Calibri"
    _BODY_PT = 9.0
    _DARK_COLOR = RGBColor(0x1A, 0x1A, 0x1A)
    _GRAY_COLOR = RGBColor(0x55, 0x55, 0x55)

    def __init__(self, output_dir: Path, data_path: Path) -> None:
        """Inicializa o gerador.
        
        Args:
            output_dir: Diretório onde os arquivos gerados serão salvos.
            data_path: Caminho para o JSON com as informações do currículo.
        """
        self._output_dir = output_dir
        self._data_path = data_path
        self._doc = Document()
        self._setup_page()

    def _setup_page(self) -> None:
        """Configura margens e fonte padrão do documento."""
        sec = self._doc.sections[0]
        sec.top_margin = sec.bottom_margin = Cm(1.1)
        sec.left_margin = sec.right_margin = Cm(1.3)

        normal = self._doc.styles["Normal"]
        normal.font.name = self._FONT_NAME
        normal.font.size = Pt(self._BODY_PT)
        normal.element.rPr.rFonts.set(qn("w:eastAsia"), self._FONT_NAME)

    @staticmethod
    def remove_spacing(par, before: float = 0, after: float = 0) -> None:
        """Zera o espaçamento do parágrafo.

        Args:
            par: Objeto parágrafo do docx.
            before: Espaçamento antes do parágrafo.
            after: Espaçamento depois do parágrafo.
        """
        par.paragraph_format.space_before = Pt(before)
        par.paragraph_format.space_after = Pt(after)
        par.paragraph_format.line_spacing = 1.0

    def add_run(
        self,
        par,
        text: str,
        *,
        bold: bool = False,
        italic: bool = False,
        size: float = _BODY_PT,
        color: RGBColor = _DARK_COLOR
    ):
        """Adiciona um texto estilizado (run) a um parágrafo."""
        run = par.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = self._FONT_NAME
        run.font.color.rgb = color
        return run

    def add_section(self, title: str) -> None:
        """Adiciona um título de seção com linha separadora."""
        par = self._doc.add_paragraph()
        self.remove_spacing(par, before=5, after=1)
        self.add_run(par, title.upper(), bold=True, size=9.5)
        
        p_pr = par._p.get_or_add_pPr()
        border = p_pr.makeelement(qn("w:pBdr"), {})
        line = p_pr.makeelement(qn("w:bottom"), {})
        line.set(qn("w:val"), "single")
        line.set(qn("w:sz"), "6")
        line.set(qn("w:space"), "1")
        line.set(qn("w:color"), "999999")
        border.append(line)
        p_pr.append(border)

    def add_bullet(self, text_parts: List[List[Any]]) -> None:
        """Adiciona um item com marcador."""
        par = self._doc.add_paragraph()
        self.remove_spacing(par, after=1)
        par.paragraph_format.left_indent = Cm(0.35)
        par.paragraph_format.first_line_indent = Cm(-0.35)
        
        self.add_run(par, "▪  ", size=self._BODY_PT)
        for txt, bold in text_parts:
            self.add_run(par, txt, bold=bold)

    def add_role(self, title: str, period: str) -> None:
        """Adiciona uma linha descritiva de cargo."""
        par = self._doc.add_paragraph()
        self.remove_spacing(par, before=3, after=1)
        self.add_run(par, title, bold=True)
        self.add_run(par, f"   {period}", italic=True, size=8.5, color=self._GRAY_COLOR)

    def convert_to_pdf(self, docx_path: Path, pdf_path: Path) -> None:
        """Converte um arquivo DOCX para PDF usando win32com."""
        try:
            import win32com.client
            logger.info("Convertendo para PDF via Word COM...")
            word = win32com.client.Dispatch("Word.Application")
            doc = word.Documents.Open(str(docx_path.resolve()))
            doc.SaveAs(str(pdf_path.resolve()), FileFormat=17)
            doc.Close()
            word.Quit()
            logger.info("PDF gerado com sucesso.")
        except Exception as e:
            logger.error("Falha ao converter para PDF: %s", e)
            raise

    def load_data(self) -> Dict[str, Any]:
        """Carrega os dados JSON do arquivo de perfil."""
        if not self._data_path.exists():
            raise FileNotFoundError(f"Arquivo de dados não encontrado: {self._data_path}")
        with open(self._data_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def build(self, filename_base: str = "curriculo___rafael_torres_nantes") -> None:
        """Constrói o documento com os dados do currículo e salva em disco."""
        data = self.load_data()

        # ---------- Cabecalho ----------
        par = self._doc.add_paragraph()
        self.remove_spacing(par)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(par, data["header"]["name"], bold=True, size=17)

        par = self._doc.add_paragraph()
        self.remove_spacing(par, after=2)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(par, data["header"]["title"], bold=True, size=9.5, color=self._GRAY_COLOR)

        par = self._doc.add_paragraph()
        self.remove_spacing(par, after=2)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(par, data["header"]["contact"], size=8.5, color=self._GRAY_COLOR)

        par = self._doc.add_paragraph()
        self.remove_spacing(par, after=1)
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.add_run(par, data["header"]["links"], size=8.5, color=self._GRAY_COLOR)

        # ---------- Perfil ----------
        self.add_section("Perfil")
        par = self._doc.add_paragraph()
        self.remove_spacing(par, after=1)
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        self.add_run(par, data["profile"]["normal_1"])
        self.add_run(par, data["profile"]["bold"], bold=True)
        self.add_run(par, data["profile"]["normal_2"])

        # ---------- Experiencia ----------
        self.add_section("Experiência")
        
        for xp in data["experience"]:
            if "grouping_title" in xp and "grouping_period" in xp:
                par = self._doc.add_paragraph()
                self.remove_spacing(par, after=1)
                self.add_run(par, xp["grouping_title"], bold=True, size=10)
                self.add_run(par, f"   {xp['grouping_period']}", italic=True, size=8.5, color=self._GRAY_COLOR)
                
            self.add_role(xp["title"], xp["period"])
            for bullet in xp["bullets"]:
                self.add_bullet(bullet)
        
        # ---------- Projetos ----------
        self.add_section("Projetos  ·  github.com/rafael-torres-nantes")
        for proj in data["projects"]:
            self.add_bullet(proj)

        # ---------- Formacao ----------
        self.add_section("Formação e Publicação")
        for edu in data["education"]:
            self.add_bullet(edu)

        # ---------- Competencias ----------
        self.add_section("Competências técnicas")
        for skill in data["skills"]:
            par = self._doc.add_paragraph()
            self.remove_spacing(par, after=1)
            self.add_run(par, f"{skill['label']}: ", bold=True)
            self.add_run(par, skill['items'])

        self._output_dir.mkdir(parents=True, exist_ok=True)
        docx_path = self._output_dir / f"{filename_base}.docx"
        pdf_path = self._output_dir / f"{filename_base}.pdf"
        
        self._doc.save(docx_path)
        logger.info("DOCX gerado com sucesso: %s", docx_path)
        
        self.convert_to_pdf(docx_path, pdf_path)


if __name__ == "__main__":
    logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
    test_dir = Path(__file__).resolve().parent.parent / "output"
    data_path = Path(__file__).resolve().parent.parent / "personal_info" / "curriculum_data.json"
    logger.info("Executando teste local de CurriculumGenerator...")
    gen = CurriculumGenerator(output_dir=test_dir, data_path=data_path)
    gen.build(filename_base="curriculo___rafael_torres_nantes_teste_local")
    logger.info("Teste concluído. Arquivo deve estar em: %s", test_dir)
